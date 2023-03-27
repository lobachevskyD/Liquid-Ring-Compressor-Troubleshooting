import tkinter as tk
from tkinter import ttk
from docx import Document
import tkinter.filedialog as filedialog
import os
import sys
import customtkinter as ctk
from os import _wrap_close
from textwrap import wrap
from PIL import Image


 # load images with light and dark mode image
image_path = os.path.join(os.path.dirname(os.path.realpath(__file__)), "images")
word_image = ctk.CTkImage(Image.open(os.path.join(image_path, "word.png")), size=(40, 40))
troubleshooting_image = ctk.CTkImage(light_image=Image.open(os.path.join(image_path, "Troubleshooting_dark.ico")),
                                                     dark_image=Image.open(os.path.join(image_path, "Troubleshooting_light.ico")), size=(40, 40))


symptom_descriptions = {
    '1': 'Low suction pressure',
    '2': 'Excessive vibration',
    '3': 'Overheating',
    '4': 'Excessive noise',
    '5': 'Reduced capacity',
    '6': 'Leakage',
    '7': 'Excessive liquid carryover',
    '8': 'Corrosion or erosion of internals',
    '9': 'Cavitation damage',
    '10': 'Seizure or failure of motor',
    '11': 'Excessive power consumption'
}

possible_causes = {
    '1': ['Insufficient liquid ring flow', 'Clogged strainer', 'Blockage in suction line', 'Low suction pressure', 'Worn impeller or casing', 'Gas leakage in system'],
    '2': ['Misalignment of motor, impeller, or coupling', 'Worn bearings', 'Unbalanced rotor', 'Excessive compressor speed'],
    '3': ['Insufficient liquid ring flow or cooling', 'Clogged strainer', 'High discharge pressure', 'Worn impeller or casing', 'Incorrect liduid level'],
    '4': ['Loose mounting bolts', 'Worn bearings', 'Damaged impeller', 'Cavitation', 'Insufficient liquid ring flow'],
    '5': ['Worn impeller or casing', 'Damaged vanes', 'Insufficient liquid ring flow', 'High discharge pressure', 'Incorrect liquid level'],
    '6': ['Worn mechanical seals', 'Insufficient flow to mechanical seal', 'Damaged gaskets or O-rings', 'Cracked casing'],
    '7': ['Insufficient liquid ring flow', 'High liquid level', 'Worn impeller or casing', 'Damaged vanes', 'High discharge pressure', 'Incorrect liquid level'],
    '8': ['Corrosive or abrasive fluid', 'Insufficient liquid ring flow', 'High temperature', 'Prolonged operation'],
    '9': ['High vacuum', 'Low liquid ring flow', 'Incorrect fluid level', 'Undersized suction piping'],
    '10': ['Overheating', 'Insufficient lubrication', 'Electrical failure', 'Incorrect voltage', 'Issues with VSD'],
    '11': ['Clogged strainer', 'Worn bearings', 'Damaged impeller or casing', 'High discharge pressure', 'Incorrect liquid level', 'Undersized or oversized motor']
}

troubleshooting_steps = {
    '1': ['Check liquid ring flow rate and clean strainer', 'Check suction pressure gauge', 'Inspect impeller and casing for wear', 'Check for blockage or gas leaks in the system'],
    '2': ['Check alignment of motor, impeller and coupling', 'Inspect bearings for wear', 'Balance rotor', 'Check pump speed against design specification'],
    '3': ['Ensure sufficient liquid ring flow', 'If fitted with heat exchanger, check thermal performance', 'Clean strainer', 'Reduce discharge pressure', 'Inspect impeller and casing for wear', 'Check liquid ring level and top up if necessary'],
    '4': ['Tighten mounting bolts', 'Inspect bearings for wear', 'Replace damaged impeller', 'Check for cavitation and adjust system accordingly', 'Check valve positions in liquid ring and bypass lines'],
    '5': ['Inspect impeller and casing for wear', 'Replace or repair impeller', 'Ensure sufficient liquid ring flow', 'Reduce discharge pressure', 'Check liquid ring level and top up if necessary'],
    '6': ['Replace mechanical seals , gaskets or O-rings and check flowrate to seal', 'Inspect casing for cracks and replace if necessary'],
    '7': ['Check liquid ring flow rate and adjust as necessary', 'Adjust liquid level', 'Inspect impeller and casing for wear', 'Replace or repair impeller', 'Reduce discharge pressure', 'Check fluid level and top up if necessary'],
    '8': ['Use appropriate materials for corrosive or abrasive fluids', 'Ensure sufficient liquid ring flow', 'Reduce operating temperature', 'Limit prolonged operation'],
    '9': ['Reduce vacuum level', 'Increase liquid ring flow', 'Check liquid ring level and top up if necessary', 'Increase suction pipe size'],
    '10': ['Investigate cause of overheating and take appropriate corrective action', 'Ensure adequate lubrication', 'Check electrical connections and replace damaged parts', 'Ensure proper voltage'],
    '11': ['Clean strainer', 'Inspect bearings for wear', 'Replace damaged impeller or casing', 'Reduce discharge pressure', 'Check liquid ring level and top up if necessary', 'Replace motor with correct size']
}

def get_result(event=None):
    #print("get_result() function called.") #troubleshooting the code - indicates calling for function in the console
    symptom_description = symptom_selector.get()
    symptom = [k for k, v in symptom_descriptions.items() if v == symptom_description][0]
    #print("Symptom selected:", symptom_description) #troubleshooting the code - console will indicate new symptom selected
    p_causes = possible_causes[symptom]
    t_steps = troubleshooting_steps[symptom]
    p_causes_text.delete("0.0", ctk.END)
    t_steps_text.delete("0.0", ctk.END)
    for cause in p_causes:
        p_causes_text.insert(ctk.END, cause + '\n')
    for step in t_steps:
        t_steps_text.insert(ctk.END, step + '\n')
    #print("Possible causes:", p_causes) #troubleshooting the code - shows possible causes as symptom is changed in the console
    #print("Troubleshooting steps:", t_steps) #troubleshooting the code - shows troubleshooting steps as symptom is changed in the console
    #print("get_result() function executed.") #troubleshooting the code - indicates execution of function

def set_default_value(event):
    default_value = symptom_selector.get()
    symptom_selector.set(default_value)

def export_to_word():
    document = Document()
    symptom_description = symptom_selector.get()
    symptom = [k for k, v in symptom_descriptions.items() if v == symptom_description][0]
    causes = possible_causes[symptom]
    steps = troubleshooting_steps[symptom]

    document.add_heading('Troubleshooting Guide', 0)
    document.add_paragraph('Symptom: ' + symptom_description)
    document.add_heading('Possible Causes:', level=1)
    for cause in causes:
        document.add_paragraph(cause)
    document.add_heading('Troubleshooting Steps:', level=1)
    for step in steps:
        document.add_paragraph(step)

    # Prompt user to select save location
    file_path = filedialog.asksaveasfilename(defaultextension='.docx', initialfile=symptom_description + ' Troubleshooting Guide')

    # Save document if file_path is not empty
    if file_path:
        document.save(file_path)

# Create GUI
root = ctk.CTk()
root.title("Liquid Ring Compressor Troubleshooting Guide")
root.resizable(False, False)  # prevent the window from being resized

frame = ctk.CTkFrame(master=root)
frame.pack(pady=20, padx=60, fill="both", expand=True)

label = ctk.CTkLabel(
    master=frame, 
    text="    Liquid Ring Compressor Troubleshooting Guide",
    image=troubleshooting_image, 
    compound="left",
    font=("Roboto", 18))
label.pack(pady=12, padx=10)


# Define purpose of the guide
description_text = "The liquid ring compressor troubleshooting guide is a comprehensive resource designed to help users diagnose and resolve issues with their compressors. The guide includes detailed symptom descriptions, a comprehensive list of possible causes for each symptom, and clear and concise troubleshooting steps for each cause. The guide is designed to be technically accurate and provides users with a clear path towards identifying and resolving issues with their compressors. By following the troubleshooting steps outlined in the guide, users can quickly and effectively diagnose and resolve issues with their liquid ring compressors, ensuring that their equipment operates smoothly and reliably over the long term."

# Add description label
text_1 = ctk.CTkTextbox(
    master=frame, 
    width=450, 
    height=100, 
    font=("Calibri", 14), 
    wrap="word")
text_1.pack(pady=15, padx=10)
text_1.insert("0.0", description_text)

# Create label for symptom selection
symptom_selector_label = ctk.CTkLabel(
    master=frame, 
    text="Select Symptom:", 
    font=("Calibri", 18))
symptom_selector_label.pack(pady=5)

# Create dropdown for symptom selection
symptom_selector = ctk.CTkComboBox(
    master=frame, 
    values=list(symptom_descriptions.values()), 
    font=("Calibri", 14), 
    state="readonly", 
    command=get_result)
#symptom_selector.bind("<<ComboboxSelected>>", get_result)
symptom_selector.bind("<Button-1>", set_default_value)
symptom_selector.pack(pady=10, padx=10)

# Create label for possible causes
p_causes_label = ctk.CTkLabel(
    master=frame, 
    text='Possible Causes:', 
    font=("Calibri", 18))
p_causes_label.pack(pady=1, padx=10)

p_causes_text = ctk.CTkTextbox(
    master=frame, 
    height=150, 
    width=500,
    corner_radius=15,
    #border_width=3,
    border_spacing=3,
    font=("Calibri", 14), 
    state="normal")
p_causes_text.pack(pady=5, padx=10)

# Create label for troubleshooting steps
t_steps_label = ctk.CTkLabel(
    master=frame, text='Troubleshooting Steps:', 
    font=("Calibri", 18))
t_steps_label.pack(pady=1, padx=10)

t_steps_text = ctk.CTkTextbox(
    master=frame, 
    height=150, 
    width=500,
    corner_radius=15,
    #border_width=3,
    border_spacing=3,
    font=("Calibri", 14), 
    state="normal")
t_steps_text.pack(pady=5, padx=10)

# Set default value for symptom selector
symptom_selector.set(list(symptom_descriptions.values())[0])

# Call get_result to display default result
get_result(None)

# Create Export to Word button
export_button = ctk.CTkButton(
    master=frame, 
    text="Export to Word",
    image=word_image, 
    compound="right",
    width=120, 
    height=40, 
    corner_radius=10, 
    font=("Calibri", 18), 
    command=export_to_word)
export_button.pack(pady=20)

#setting GUI appearance mode
def change_appearance_mode_event(new_appearance_mode):
    ctk.set_appearance_mode(new_appearance_mode)

appearance_mode_menu = ctk.CTkOptionMenu(
    master=frame, 
    values=["Light", "Dark", "System"],
    command=change_appearance_mode_event)
appearance_mode_menu.pack(padx=20, pady=20)

# Create exit button
exit_button = ctk.CTkButton(
    master=frame, 
    text="Exit",
    fg_color="red",
    hover_color="dark red",
    font=("Calibri", 18), 
    command=root.destroy)
exit_button.pack(pady=20)


root.mainloop()


